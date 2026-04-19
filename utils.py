"""
工具函数模块
"""

from io import BytesIO

from docx import Document


def format_output(text) -> str:
    """格式化输出文本，确俜Markdown格式正确"""
    # 处理None或空值
    if text is None:
        return "暂无内容"

    # 处理列表类型
    if isinstance(text, list):
        items = [str(item).strip() for item in text if str(item).strip()]
        if not items:
            return "暂无内容"
        text = "\n".join(f"- {item}" for item in items)
    
    # 处理字典类型（如风险分析结果）
    if isinstance(text, dict):
        # 如果有raw_analysis字段，使用它
        if 'raw_analysis' in text and text['raw_analysis']:
            text = text['raw_analysis']
        else:
            # 否则将字典转换为字符串
            import json
            text = json.dumps(text, ensure_ascii=False, indent=2)
    
    # 再次检查是否为空
    if not text:
        return "暂无内容"
    
    # 确俜是字符串类型
    if not isinstance(text, str):
        text = str(text)
    
    # 清理多余的空行
    lines = text.split('\n')
    cleaned_lines = []
    prev_empty = False
    
    for line in lines:
        is_empty = line.strip() == ''
        if is_empty and prev_empty:
            continue
        cleaned_lines.append(line)
        prev_empty = is_empty
    
    return '\n'.join(cleaned_lines)


def create_example_cases() -> dict:
    """创建示例需求案例"""
    return {
        "案例A-信息不足：会议纪要AI想法（会被拦截）": """我们想做一个会议纪要 AI 功能。

大概想法：
1. 用户上传录音后，系统自动生成纪要
2. 能提炼重点、输出待办事项
3. 最好还能给老板看一个简版总结

现在还没想清楚：
- 先给哪些用户用
- 支持什么文件格式
- 输出成什么样
- 出错时怎么处理
- 上线后看什么指标

希望先看看总控Agent会提示我们还缺哪些信息。""",

        "案例B-信息完整：智能客服系统升级（可进入转译）": """我们计划为现有SaaS产品升级一套智能客服系统，目标是在不明显降低用户满意度的前提下，减少人工客服重复答疑压力，提高夜间和高峰时段的自动接待能力。

目标用户：
1. SaaS产品的付费企业管理员
2. 企业员工端普通使用者
3. 内部客服团队和运营团队

核心场景：
1. 用户在帮助中心、产品内消息窗口发起问题咨询
2. 夜间或人工排队时优先由AI客服接待
3. 用户查询账号配置、权限设置、计费规则、常见报错处理方法
4. AI无法解决时无缝转人工，并带上上下文和问题摘要

核心功能：
1. 7x24小时文本客服接待
2. 基于FAQ、产品文档、工单历史进行回答
3. 支持多轮对话，保留本轮会话上下文
4. 高风险或高情绪问题自动转人工
5. 给客服后台输出问题分类、转人工原因、解决率等分析数据

成功标准：
1. 自动化解决率达到65%以上
2. 用户满意度不低于4.5/5
3. 平均首次响应时间小于3秒
4. 人工客服重复问题工单量下降30%

约束条件：
1. 第一阶段只支持中文文本，不支持语音和图片
2. 只能接入现有帮助中心知识库和近12个月工单数据
3. 涉及退款、合同、账户封禁的问题必须转人工
4. MVP版本需在6周内上线，先对20%的企业租户灰度

依赖与接口：
1. 依赖帮助中心知识库接口
2. 依赖工单系统读取历史问答
3. 依赖用户中心提供用户身份、套餐、企业信息
4. 依赖客服工作台接收转人工会话

AI范围与兜底：
1. AI负责问题理解、答案生成、情绪识别、转人工建议
2. 命中高风险规则时直接转人工
3. AI置信度低于阈值时不给确定答案，只给建议并转人工

测试与验收要求：
1. 需要覆盖正常咨询、连续追问、越权提问、敏感问题、低置信度回答等场景
2. 要明确记录什么情况算可接受回答，什么情况算错误回答

运营要求：
1. 对外卖点是“更快响应、夜间可用、复杂问题无缝转人工”
2. 上线后重点看自动化解决率、转人工率、满意度、留存投诉量

风险与待确认：
1. 需要法务确认工单历史数据是否可用于模型检索
2. 需要客服团队确认高风险问题清单
3. 需要产品确定是否允许AI直接回答计费类问题""",

        "案例C-信息完整：个性化推荐引擎2.0（可进入转译）": """我们要为内容平台升级个性化推荐引擎2.0，目标是在保证内容多样性和平台合规的前提下，提高首页推荐点击率、用户停留时长和次日留存。

目标用户：
1. 阅读图文和短视频内容的普通用户
2. 刚注册的新用户
3. 内容创作者和运营团队（间接受影响）

核心场景：
1. 用户打开首页信息流时看到个性化推荐内容
2. 用户浏览某篇内容后实时得到相关推荐
3. 新用户首次进入时触发冷启动推荐
4. 运营团队查看不同推荐策略的效果对比

核心功能：
1. 基于浏览、点赞、收藏、分享、评论行为建模
2. 同时考虑内容标签、分类、时效性、热度和质量分
3. 首页离线推荐列表每日更新
4. 详情页实时相关推荐接口毫秒级返回
5. 新用户、新内容冷启动策略
6. 多样性控制和低质内容过滤
7. A/B测试框架支持策略切换
8. 提供“为什么推荐这条内容”的可解释文案

成功标准：
1. 首页推荐点击率提升20%
2. 用户平均停留时长提升15%
3. 次日留存提升8%
4. 多样性指标不低于当前基线
5. 负反馈率不上升

约束条件：
1. 首页推荐接口P95响应时间必须小于100ms
2. 第一阶段只覆盖图文与短视频，不覆盖直播
3. 推荐结果必须满足平台内容审核和未成年人保护规则
4. 算法训练与推理成本需要控制在现有预算上限内

依赖与接口：
1. 依赖用户行为埋点系统
2. 依赖内容标签和审核结果数据
3. 依赖实验平台做A/B测试分桶
4. 依赖推荐服务接口为首页和详情页供给结果

AI/算法范围：
1. 召回、粗排、精排由推荐算法负责
2. 可解释文案由规则+轻量生成模型共同完成
3. 新用户冷启动要结合热门内容、类目偏好问答和地域特征

测试与验收要求：
1. 需要覆盖新用户、老用户、低活跃用户、高活跃用户等不同分群
2. 需要验证实时推荐、离线推荐、冷启动、多样性控制、解释文案准确性
3. 需要验证高并发、缓存失效、数据延迟等异常情况

运营要求：
1. 需要给运营后台展示不同策略版本的点击率、停留时长、留存变化
2. 需要支持活动期人工干预部分内容曝光

风险与待确认：
1. 需要确认解释文案是否会误导用户理解推荐逻辑
2. 需要确认活动流量和自然流量混排规则
3. 需要确认多样性指标与商业化目标之间的优先级""",
    }


def extract_text_from_docx(uploaded_file) -> str:
    """从上传的Word文档中提取纯文本内容，仅支持.docx"""
    file_bytes = uploaded_file.getvalue()
    document = Document(BytesIO(file_bytes))

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs).strip()


def save_to_markdown(filename: str, content: str):
    """保存内容为Markdown文件"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(content)


def calculate_token_usage(text: str) -> int:
    """估算token数量（粗略估算：1个中文约1.5个token，1个英文单词约1.3个token）"""
    chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
    english_words = len(text.split())
    return int(chinese_chars * 1.5 + english_words * 1.3)
